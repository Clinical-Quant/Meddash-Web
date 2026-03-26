"""
export_to_docx.py — Meddash Markdown to Word Converter
======================================================
Converts the generated Meddash Markdown reports (KOL Briefs, TA Landscapes)
into professional Word documents (.docx), complying with the consulting-grade
Meddash Intelligence Style Guide.
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_header(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    
    # Apply strict design system from Style Guide
    if level == 1:
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 60, 93) # #0B3C5D Dark Blue
    elif level == 2:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 60, 93) # #0B3C5D Dark Blue
    elif level == 3:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(50, 140, 193) # #328CC1 Muted Blue
    return p

def process_table(doc, lines, start_idx):
    headers = [col.strip() for col in lines[start_idx].strip('|').split('|') if col.strip()]
    
    rows = []
    idx = start_idx + 2 
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        row_data = [col.strip() for col in line.strip('|').split('|')]
        rows.append(row_data)
        idx += 1
        
    if not headers or not rows:
        return idx
        
    table = doc.add_table(rows=1, cols=len(headers))
    # 'Light Shading Accent 1' provides a blue header and alternating rows in default Word templates
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, hc in enumerate(headers):
        if i < len(hdr_cells):
            run = hdr_cells[i].paragraphs[0].add_run(hc)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = True
            
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            if i < len(row_cells):
                # Clean up KOL signals for Word
                clean_val = val.replace('⚪', '•').replace('⭐', '*').replace('🟡', '-').replace('🔵', '.')
                para = row_cells[i].paragraphs[0]
                run = para.add_run(clean_val)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                
    doc.add_paragraph()
    return idx

def convert_md_to_docx(md_path, out_path=None):
    if not os.path.exists(md_path):
        print(f"Error: File not found at {md_path}")
        sys.exit(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    doc = Document()
    
    # 1. Page Layout (A4/US Letter, 1 inch margins)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    
    # 2. Header and Footer Configuration
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Meddash Intelligence\t\t\tTherapeutic Area Landscape Report"
    hp.style.font.name = 'Calibri'
    hp.style.font.size = Pt(9)
    # Using tabs to push right side
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential — Prepared for [Client Name]"
    fp.style.font.name = 'Calibri'
    fp.style.font.size = Pt(9)
    
    # 3. Font System (Primary: Calibri 11)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    idx = 0
    
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Headers
        if line.startswith('# '):
            apply_header(doc, line[2:].replace('**', '').strip(), 1)
        elif line.startswith('## '):
            apply_header(doc, line[3:].replace('**', '').strip(), 2)
        elif line.startswith('### '):
            apply_header(doc, line[4:].replace('**', '').strip(), 3)
            
        # Images Let's assume standard md format: ![alt](path)
        elif re.match(r'!\[.*?\]\((.*?)\)', line):
            match = re.match(r'!\[.*?\]\((.*?)\)', line)
            img_path = match.group(1)
            
            # Resolve relative path considering the md file location
            base_dir = os.path.dirname(os.path.abspath(md_path))
            full_img_path = os.path.join(base_dir, img_path)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(full_img_path):
                r = p.add_run()
                r.add_picture(full_img_path, width=Inches(6.0))
                # Add caption
                alt_text = line[line.find("[")+1:line.find("]")]
                caption = doc.add_paragraph(alt_text)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.name = 'Calibri'
                caption.runs[0].font.size = Pt(10)
                caption.runs[0].font.italic = True
                caption.runs[0].font.color.rgb = RGBColor(85, 85, 85) # #555555
            else:
                p.add_run(f"[Image not found: {img_path}]").italic = True
                
        # Tables
        elif line.startswith('|') and idx + 1 < len(lines) and lines[idx+1].strip().startswith('|-'):
            idx = process_table(doc, lines, idx)
            continue
            
        # Blockquotes / Notes
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:], style='Intense Quote')
            
        # Bullet points
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
            
        # Unordered Lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).replace('**', '')
            doc.add_paragraph(text, style='List Number')
            
        # Horizontal rules
        elif line == '---':
            doc.add_page_break() # Convert --- to page breaks for cleaner reporting
            
        # Bold strings at start like **Indication:** Multiple Myeloma
        elif line.startswith('**') and ':**' in line:
            parts = line.split(':**', 1)
            if len(parts) == 2:
                p = doc.add_paragraph()
                r1 = p.add_run(parts[0].replace('**', '') + ':')
                r1.bold = True
                r2 = p.add_run(parts[1])
            else:
                doc.add_paragraph(line.replace('**', ''))
        
        # Normal text
        elif line:
            clean_line = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_line)
            
        idx += 1

    if out_path is None:
        out_name = os.path.splitext(md_path)[0] + '.docx'
    else:
        out_name = out_path
        
    doc.save(out_name)
    print(f"✅ Exported to Word: {out_name}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python export_to_docx.py <markdown_file> [optional_output_file]")
        sys.exit(1)
        
    out_file = sys.argv[2] if len(sys.argv) > 2 else None
    convert_md_to_docx(sys.argv[1], out_path=out_file)
